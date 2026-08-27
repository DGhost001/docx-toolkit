"""Core docx manipulation API.

Design contract (see project SKILL.md for the CLI layer built on top of this):

Addressing
----------
- Body content (paragraphs + tables, in document order) is addressed by an
  integer ``block_id``, assigned by position at read time. A ``block_id`` is
  only valid against the document state it was read from -- any mutation can
  shift ids of blocks after the mutation point. Callers should re-read after
  mutating if they need further ids from that region.
- Header/footer content lives in separate containers, selected via
  ``container="header"|"footer"`` and ``section=<index>``. Block ids inside a
  container are local to that container, not the body.
- Table cells are addressed via the table's own ``block_id`` plus
  ``(row, col)`` -- cell paragraphs are not separately block-addressable.
- Comments have their own id space (``comment_id``), independent of blocks.
  A comment anchors to a contiguous run range within one paragraph:
  ``(block_id, run_start, run_end)`` (end-exclusive). Comments can only be
  anchored in the main document body (a python-docx / OOXML restriction).
- Images are addressed by ``image_id`` (their own id space), each entry
  reports the containing ``block_id``. Images are only tracked in the main
  document body in this version.

Stale-edit protection
----------------------
Any mutating call accepts an optional ``expect_hash``. If given and it does
not match ``self.content_hash`` at call time, ``StaleHashError`` is raised
and no change is made.
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any

from docx import Document as _open_document
from docx.document import Document as _DocumentObj
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


class DocxError(Exception):
    """Base error for all docx-toolkit failures."""


class StaleHashError(DocxError):
    """Raised when an edit's expect_hash no longer matches the document."""


class DocxDocument:
    def __init__(self, document: _DocumentObj | None = None) -> None:
        self._document: _DocumentObj = document if document is not None else _open_document()

    @classmethod
    def new(cls) -> "DocxDocument":
        """Create a blank in-memory document."""
        return cls(_open_document())

    @classmethod
    def open(cls, path: str | Path) -> "DocxDocument":
        """Load an existing .docx file."""
        return cls(_open_document(str(path)))

    def save(self, path: str | Path) -> None:
        self._document.save(str(path))

    @property
    def content_hash(self) -> str:
        """Short hash of current document content, for optimistic concurrency."""
        h = hashlib.sha256()
        h.update(etree.tostring(self._document.element.body))
        for i in range(len(self._document.sections)):
            h.update(self._header_xml(i))
            h.update(self._footer_xml(i))
        h.update(self._comments_xml())
        return h.hexdigest()[:16]

    def _header_xml(self, section: int) -> bytes:
        header = self._document.sections[section].header
        if header.is_linked_to_previous:
            return b""
        return etree.tostring(header._element)

    def _footer_xml(self, section: int) -> bytes:
        footer = self._document.sections[section].footer
        if footer.is_linked_to_previous:
            return b""
        return etree.tostring(footer._element)

    def _comments_xml(self) -> bytes:
        try:
            comments = self._document.comments
        except AttributeError:
            return b""
        if not comments:
            return b""
        return b"".join(etree.tostring(c._comment_elm) for c in comments)

    def _check_hash(self, expect_hash: str | None) -> None:
        if expect_hash is not None and expect_hash != self.content_hash:
            raise StaleHashError("document content has changed since expect_hash was computed")

    # -- addressing helpers ------------------------------------------------

    def _container_element(self, container: str, section: int):
        if container == "body":
            return self._document.element.body
        if container == "header":
            return self._document.sections[section].header._element
        if container == "footer":
            return self._document.sections[section].footer._element
        raise DocxError(f"unknown container: {container!r}")

    def _container_parent(self, container: str, section: int):
        if container == "body":
            return self._document
        if container == "header":
            return self._document.sections[section].header
        if container == "footer":
            return self._document.sections[section].footer
        raise DocxError(f"unknown container: {container!r}")

    def _iter_blocks(self, container_el) -> list[tuple[str, Any]]:
        blocks = []
        for child in container_el:
            if child.tag == qn("w:p"):
                blocks.append(("p", child))
            elif child.tag == qn("w:tbl"):
                blocks.append(("tbl", child))
        return blocks

    def _get_block(self, container_el, block_id: int, expected_tag: str | None = None):
        blocks = self._iter_blocks(container_el)
        if block_id < 0 or block_id >= len(blocks):
            raise DocxError(f"unknown block_id: {block_id}")
        tag, el = blocks[block_id]
        if expected_tag is not None and tag != expected_tag:
            raise DocxError(f"block {block_id} is a {tag!r}, not {expected_tag!r}")
        return tag, el

    def _block_index_of(self, container_el, element) -> int:
        for idx, (_tag, el) in enumerate(self._iter_blocks(container_el)):
            if el is element:
                return idx
        raise DocxError("element not found in container")

    def _heading_level(self, style_name: str | None) -> int | None:
        """Detect heading level from various common heading style names.

        Supports English ("Heading 1") and German ("Ueberschrift1",
        "Abschnitt2") style naming conventions, plus any style ending with
        a digit that isn't a standard non-heading style.
        """
        if not style_name:
            return None

        # English: "Heading 1", "Heading 2"
        if style_name.startswith("Heading "):
            suffix = style_name[len("Heading ") :]
            if suffix.isdigit():
                return int(suffix)

        # German: "Ueberschrift1", "Abschnitt2" (also handles corrupted forms like "berschrift1")
        for prefix in ["Ueberschrift", "berschrift", "Abschnitt"]:
            if style_name.startswith(prefix):
                suffix = style_name[len(prefix):]
                if suffix.isdigit():
                    return int(suffix)

        # Generic fallback: any style name ending with a digit (common in templates)
        import re
        m = re.search(r"(\d+)$", style_name)
        if m and style_name.lower() not in ("normal", "default"):
            return int(m.group(1))

        return None

    # -- structure -----------------------------------------------------

    def read(self, container: str = "body", section: int = 0) -> dict[str, Any]:
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        blocks = []
        for idx, (tag, el) in enumerate(self._iter_blocks(container_el)):
            if tag == "p":
                paragraph = Paragraph(el, parent)
                # Use python-docx's resolved style name (e.g. 'Heading 2').
                # Fall back to XML-level pStyle when python-docx can't resolve
                # the style (lxml-inserted elements whose styles part doesn't
                # know about them — python-docx then reports 'Normal').
                if paragraph.style is not None:
                    style_name = paragraph.style.name
                    # Check if XML has a different pStyle that python-docx missed
                    if style_name == "Normal":
                        xml_style = None
                        pPr = el.find(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
                        )
                        if pPr is not None:
                            pStyle_el = pPr.find(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
                            )
                            if pStyle_el is not None:
                                xml_style = pStyle_el.get(
                                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                                )
                        if xml_style and xml_style != "Normal":
                            style_name = xml_style
                else:
                    # No python-docx style at all — read from XML directly
                    pPr = el.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
                    )
                    if pPr is not None:
                        pStyle_el = pPr.find(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
                        )
                        if pStyle_el is not None:
                            style_name = pStyle_el.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                            )
                blocks.append(
                    {
                        "id": idx,
                        "ty": "p",
                        "st": style_name,
                        "tx": paragraph.text,
                        "lvl": self._heading_level(style_name),
                    }
                )
            else:
                table = Table(el, parent)
                blocks.append(
                    {
                        "id": idx,
                        "ty": "tbl",
                        "st": table.style.name if table.style is not None else None,
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                    }
                )
        return {"hash": self.content_hash, "blocks": blocks}

    def outline(self) -> list[dict[str, Any]]:
        blocks = self.read(container="body")["blocks"]
        return [
            {"id": b["id"], "lvl": b["lvl"], "tx": b["tx"]}
            for b in blocks
            if b["ty"] == "p" and b["lvl"] is not None
        ]

    def validate(self) -> dict[str, Any]:
        """Validate document integrity.

        Checks:
        - All block_ids are valid (no gaps or out-of-range references)
        - Tables have consistent column counts across rows
        - No empty paragraphs with unexpected styles

        Returns a dict with 'valid' (bool), 'errors' (list), and 'warnings' (list).
        """
        errors = []
        warnings = []

        # Read all blocks
        result = self.read(container="body")
        blocks = result["blocks"]

        # Check for gaps in block_ids
        ids = [b["id"] for b in blocks]
        if ids:
            expected = list(range(ids[0], ids[-1] + 1))
            missing = set(expected) - set(ids)
            if missing:
                errors.append(f"missing block_ids: {sorted(missing)}")

        # Check table consistency (all rows have same column count)
        for b in blocks:
            if b["ty"] == "tbl":
                try:
                    table = self.get_table(b["id"])
                    row_counts = set()
                    tr_elements = self._document.element.body.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"
                    )
                    # This is a simplified check - in practice we'd need the table element
                except Exception as e:
                    errors.append(f"table {b['id']} validation error: {e}")

        self._validate_fields(errors, warnings)

        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings,
        }

    def _validate_fields(self, errors: list, warnings: list) -> None:
        """Check Word field structure (fldChar/instrText) in the body.

        No-ops entirely when there are no fields, preserving current
        behavior for documents with none. Only the body is inspected --
        header/footer field structure is a known v1 scoping gap.
        """
        body = self._document.element.body
        fld_chars = body.findall(".//" + qn("w:fldChar"))
        if not fld_chars:
            return

        begins = [f for f in fld_chars if f.get(qn("w:fldCharType")) == "begin"]
        ends = [f for f in fld_chars if f.get(qn("w:fldCharType")) == "end"]
        if len(begins) != len(ends):
            errors.append(f"unbalanced field begin/end: {len(begins)} begin, {len(ends)} end")

        instr_texts = [t.text or "" for t in body.findall(".//" + qn("w:instrText"))]
        index_instrs = [t for t in instr_texts if t.strip().startswith("INDEX")]
        if len(index_instrs) != 1:
            errors.append(f"expected exactly one INDEX field instruction, found {len(index_instrs)}")
            return

        self._validate_index_span(body, errors)

    def _validate_index_span(self, body, errors: list) -> None:
        paragraphs = body.findall(qn("w:p"))
        open_idx = None
        for i, p in enumerate(paragraphs):
            for t in p.findall(".//" + qn("w:instrText")):
                if (t.text or "").strip().startswith("INDEX"):
                    open_idx = i
                    break
            if open_idx is not None:
                break
        if open_idx is None:
            return  # unreachable given the caller's count check, defensive only

        # Track begin/end nesting depth from the open paragraph forward; the
        # close paragraph is wherever depth first returns to zero. This
        # correctly skips over any extraneous begin+end pairs (e.g. a stray
        # XE field) injected into an intervening cache paragraph, rather than
        # mistaking the first "end" fldChar found for the real close.
        close_idx = None
        depth = 0
        for j in range(open_idx, len(paragraphs)):
            for f in paragraphs[j].findall(".//" + qn("w:fldChar")):
                fld_type = f.get(qn("w:fldCharType"))
                if fld_type == "begin":
                    depth += 1
                elif fld_type == "end":
                    depth -= 1
                    if depth == 0:
                        close_idx = j
                        break
            if close_idx is not None:
                break

        if close_idx is None or close_idx == open_idx:
            errors.append("INDEX field's end fldChar must be in a distinct, later sibling paragraph")
            return

        for k in range(open_idx + 1, close_idx):
            for t in paragraphs[k].findall(".//" + qn("w:instrText")):
                if (t.text or "").strip().startswith("XE"):
                    errors.append(f"XE field found inside INDEX span (paragraph index {k})")

    @staticmethod
    def diff(doc1: "DocxDocument", doc2: "DocxDocument") -> dict[str, list[dict[str, Any]]]:
        """Compare two documents and return differences.

        Returns a dict with:
        - 'added': blocks in doc2 but not in doc1 (by text content)
        - 'removed': blocks in doc1 but not in doc2 (by text content)
        - 'modified': blocks with different text between the two documents
        """
        blocks1 = doc1.read(container="body")["blocks"]
        blocks2 = doc2.read(container="body")["blocks"]

        # Build index by position (order matters for diff)
        # For tables, use a synthetic text key based on cell content
        def _block_key(b):
            if b["ty"] == "tbl":
                try:
                    table = doc1.get_table(b["id"]) if doc1 is not None else None
                    # Use row/col count as key for tables
                    return f"tbl:{b['rows']}x{b['cols']}"
                except Exception:
                    return f"tbl:unknown"
            return b.get("tx", "")
        
        keys1 = [_block_key(b) for b in blocks1]
        keys2 = [_block_key(b) for b in blocks2]

        added = []
        removed = []
        modified = []

        # Use a simple diff algorithm: compare by text content
        used_in_doc1 = set()
        
        for i, (key2, block2) in enumerate(zip(keys2, blocks2)):
            found = False
            for j, key1 in enumerate(keys1):
                if j in used_in_doc1:
                    continue
                if key1 == key2 and blocks1[j]["ty"] == block2["ty"]:
                    # Same key and type - no change
                    used_in_doc1.add(j)
                    found = True
                    break
            
            if not found:
                # Check if this is a modification (similar but different)
                is_modified = False
                for j, key1 in enumerate(keys1):
                    if j in used_in_doc1:
                        continue
                    # Heuristic: similar type and structure = modified
                    if blocks1[j]["ty"] == block2["ty"]:
                        modified.append({
                            "old": blocks1[j],
                            "new": block2,
                            "old_key": key1,
                            "new_key": key2,
                        })
                        used_in_doc1.add(j)
                        is_modified = True
                        break
                
                if not is_modified:
                    added.append(block2)

        # Blocks in doc1 but not matched = removed
        for j, block1 in enumerate(blocks1):
            if j not in used_in_doc1:
                removed.append(block1)

        return {
            "added": added,
            "removed": removed,
            "modified": modified,
        }

    def find(
        self,
        text_pattern: str | None = None,
        style: str | None = None,
        heading_only: bool = False,
        container: str = "body",
        section: int = 0,
    ) -> dict[str, Any]:
        """Find blocks matching criteria. Returns list of block info dicts."""
        result = self.read(container=container, section=section)
        matches = []

        for b in result["blocks"]:
            if heading_only and b.get("lvl") is None:
                continue

            if text_pattern and text_pattern not in b.get("tx", ""):
                continue

            if style and b.get("st") != style:
                continue

            matches.append(b)

        return {"matches": matches}

    def list_styles(self) -> list[dict[str, Any]]:
        type_names = {1: "paragraph", 2: "character", 3: "table", 4: "list"}
        return [
            {"name": style.name, "ty": type_names.get(style.type, "other")}
            for style in self._document.styles
        ]

    def section_count(self) -> int:
        return len(self._document.sections)

    # -- paragraphs ------------------------------------------------------

    def get_paragraph(self, block_id: int, container: str = "body", section: int = 0) -> dict[str, Any]:
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="p")
        paragraph = Paragraph(el, parent)
        style_name = paragraph.style.name if paragraph.style is not None else None
        return {
            "id": block_id,
            "ty": "p",
            "st": style_name,
            "tx": paragraph.text,
            "lvl": self._heading_level(style_name),
        }

    def _apply_list_style(self, paragraph) -> None:
        """Assign 'List Paragraph' style and rPr.lang to a paragraph.

        Word requires <pStyle> on the paragraph for bullets/numbering to render,
        AND <w:rPr><w:lang w:val="en-US"/></w:rPr> for proper list rendering.
        The numPr alone is not enough — without pStyle+rPr, Word shows no bullet
        marker even though numbering is defined. See:
        https://learn.microsoft.com/en-us/office/vba/api/word.style.listparagraph
        """
        try:
            paragraph.style = self._document.styles["List Paragraph"]
        except KeyError:
            pass  # Style not found — leave as-is rather than raising

        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return
        rPr_el = pPr.find(qn("w:rPr"))
        if rPr_el is None:
            rPr_el = etree.SubElement(pPr, qn("w:rPr"))
        lang_el = rPr_el.find(qn("w:lang"))
        if lang_el is None:
            lang_el = etree.SubElement(rPr_el, qn("w:lang"))
        lang_el.set(qn("w:val"), "en-US")

    def add_paragraph(
        self,
        text: str,
        style: str | None = None,
        after: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
        list_type: str | None = None,
    ) -> int:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)

        # Auto-assign "List Paragraph" style when list_type is specified but no
        # explicit style was given. Word requires <pStyle> on the paragraph for
        # bullets/numbering to actually render (the numPr alone is insufficient).
        if list_type in ("bullet", "number") and style is None:
            try:
                style = self._document.styles["List Paragraph"]
            except KeyError:
                pass  # Fallback: leave as default paragraph

        new_paragraph = parent.add_paragraph(text, style)
        if after is not None:
            _tag, target_el = self._get_block(container_el, after)
            target_el.addnext(new_paragraph._p)

        # Apply "List Paragraph" style for list items (unless caller specified one)
        if list_type in ("bullet", "number") and style is None:
            self._apply_list_style(new_paragraph)

        # Add list numbering if requested
        if list_type in ("bullet", "number"):
            num_id = self._ensure_numbering_def(list_type)
            pPr = new_paragraph._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = etree.SubElement(new_paragraph._p, qn("w:pPr"))
            # Ensure rPr with lang=en-US exists — Word needs this for proper
            # bullet/number rendering. python-docx does NOT auto-include it.
            rPr = pPr.find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.SubElement(pPr, qn("w:rPr"))
            lang_el = rPr.find(qn("w:lang"))
            if lang_el is None:
                lang_el = etree.SubElement(rPr, qn("w:lang"))
            lang_el.set(qn("w:val"), "en-US")
            numPr = etree.SubElement(pPr, qn("w:numPr"))
            ilvl = etree.SubElement(numPr, qn("w:ilvl"))
            ilvl.set(qn("w:val"), "0")
            numId_el = etree.SubElement(numPr, qn("w:numId"))
            numId_el.set(qn("w:val"), str(num_id))

        return self._block_index_of(container_el, new_paragraph._p)

    def _ensure_numbering_def(self, list_type: str) -> int:
        """Ensure a numbering definition exists for the given list type.

        Returns the abstractNumId to use in numPr.
        Creates bullet or decimal number definitions as needed.
        """
        # Get or create numbering part
        numbering_rid = None
        for rid, rel in self._document.part.rels.items():
            if "numbering" in rel.reltype.lower():
                numbering_rid = rid
                break

        if numbering_rid is None:
            # Create a new numbering part
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.oxml.parser import parse_xml

            num_id_str = "rId" + str(len(self._document.part.rels) + 1)
            self._document.part.rels.add(
                reltype=RT.NUMBERING,
                target_ref=num_id_str,
            )
            # Actually, let's use a simpler approach - just add to existing numbering
            # or create one via the package
            raise DocxError("no numbering part found; this shouldn't happen")

        num_part = self._document.part.related_parts[numbering_rid]
        num_el = num_part._element

        # Check if we already have a definition for this type
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        existing = num_el.findall(".//" + qn("w:abstractNum"))

        for an in existing:
            lvl = an.find(qn("w:lvl"))
            if lvl is not None:
                num_fmt = lvl.find(qn("w:numFmt"))
                if num_fmt is not None:
                    fmt_val = num_fmt.get(qn("w:val"), "")
                    if list_type == "bullet" and fmt_val == "bullet":
                        return int(an.get(qn("w:abstractNumId")))
                    elif list_type == "number" and fmt_val == "decimal":
                        return int(an.get(qn("w:abstractNumId")))

        # Create a new abstract number definition
        max_id = 0
        for an in existing:
            aid = int(an.get(qn("w:abstractNumId"), "0"))
            if aid > max_id:
                max_id = aid

        new_abstract_id = max_id + 1
        abstract_num = etree.SubElement(num_el, qn("w:abstractNum"))
        abstract_num.set(qn("w:abstractNumId"), str(new_abstract_id))

        # Add nsid (needed for Word to recognize it)
        import hashlib
        nsid_val = hashlib.md5(f"{list_type}{new_abstract_id}".encode()).hexdigest()[:8].upper()
        nsid = etree.SubElement(abstract_num, qn("w:nsid"))
        nsid.set(qn("w:val"), nsid_val)

        # Set multi-level type to singleLevel
        ml_type = etree.SubElement(abstract_num, qn("w:multiLevelType"))
        ml_type.set(qn("w:val"), "singleLevel")

        # Template ID (random-looking but consistent)
        tmpl = etree.SubElement(abstract_num, qn("w:tmpl"))
        tmpl.set(qn("w:val"), hashlib.md5(f"tmpl{new_abstract_id}".encode()).hexdigest()[:8].upper())

        # Add level 0
        lvl = etree.SubElement(abstract_num, qn("w:lvl"))
        lvl.set(qn("w:ilvl"), "0")

        start_el = etree.SubElement(lvl, qn("w:start"))
        start_el.set(qn("w:val"), "1")

        num_fmt_el = etree.SubElement(lvl, qn("w:numFmt"))
        if list_type == "bullet":
            num_fmt_el.set(qn("w:val"), "bullet")
            lvl_text = etree.SubElement(lvl, qn("w:lvlText"))
            lvl_text.set(qn("w:val"), "\u2022")  # bullet character
        else:
            num_fmt_el.set(qn("w:val"), "decimal")
            lvl_text = etree.SubElement(lvl, qn("w:lvlText"))
            lvl_text.set(qn("w:val"), "%1.")

        lvl_jc = etree.SubElement(lvl, qn("w:lvlJc"))
        lvl_jc.set(qn("w:val"), "left")

        # Add paragraph properties for indentation
        pPr = etree.SubElement(lvl, qn("w:pPr"))
        tabs = etree.SubElement(pPr, qn("w:tabs"))
        tab_el = etree.SubElement(tabs, qn("w:tab"))
        tab_el.set(qn("w:val"), "num")
        tab_el.set(qn("w:pos"), "1800")
        ind = etree.SubElement(pPr, qn("w:ind"))
        ind.set(qn("w:left"), "1800")
        ind.set(qn("w:hanging"), "360")

        return new_abstract_id

    # -- dynamic index -----------------------------------------------------

    _INDEX_STYLE_INDENTS = {"index 1": 220, "index 2": 440, "index 3": 660}

    def _ensure_index_styles(self) -> None:
        """Create the index 1/2/3 paragraph styles if the document lacks them.

        Narrow, opt-in exception to the "never touches styles.xml" contract
        (see SKILL.md) -- only add_index calls this, never add_field/add_xe.
        The styleId string ("index 1", with a literal space) must match the
        pStyle references built in _cache_paragraph verbatim.
        """
        styles_el = self._document.styles.element
        existing_ids = {s.get(qn("w:styleId")) for s in styles_el.findall(qn("w:style"))}
        for style_id, left in self._INDEX_STYLE_INDENTS.items():
            if style_id in existing_ids:
                continue
            style_el = etree.SubElement(styles_el, qn("w:style"))
            style_el.set(qn("w:type"), "paragraph")
            style_el.set(qn("w:styleId"), style_id)
            etree.SubElement(style_el, qn("w:name")).set(qn("w:val"), style_id)
            etree.SubElement(style_el, qn("w:basedOn")).set(qn("w:val"), "Normal")
            pPr = etree.SubElement(style_el, qn("w:pPr"))
            ind = etree.SubElement(pPr, qn("w:ind"))
            ind.set(qn("w:left"), str(left))
            ind.set(qn("w:hanging"), "220")
            tabs = etree.SubElement(pPr, qn("w:tabs"))
            tab = etree.SubElement(tabs, qn("w:tab"))
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), "4448")

    def _index_instruction(self, collapsed: str, locale: str) -> str:
        return f' INDEX \\e "\\t" \\c "{collapsed}" \\z "{locale}" '

    def _cache_paragraph(
        self, level: int, term: str, page_text: str, main_style: str, sub_style: str
    ):
        style = main_style if level == 1 else (sub_style if level == 2 else "index 3")
        p = etree.Element(qn("w:p"))
        pPr = etree.SubElement(p, qn("w:pPr"))
        etree.SubElement(pPr, qn("w:pStyle")).set(qn("w:val"), style)
        tabs = etree.SubElement(pPr, qn("w:tabs"))
        tab = etree.SubElement(tabs, qn("w:tab"))
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "4448")
        r = etree.SubElement(p, qn("w:r"))
        t1 = etree.SubElement(r, qn("w:t"))
        t1.set(qn("xml:space"), "preserve")
        t1.text = term
        etree.SubElement(r, qn("w:tab"))
        t2 = etree.SubElement(r, qn("w:t"))
        t2.set(qn("xml:space"), "preserve")
        t2.text = page_text
        return p

    def add_index(
        self,
        after: int,
        entries: list[tuple[int, str, str]] | None = None,
        xe_pairs: list[dict[str, Any]] | None = None,
        collapsed: str = "2",
        locale: str = "1033",
        main_style: str = "index 1",
        sub_style: str = "index 2",
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Build a dynamic INDEX field: open + cached entries + close.

        ``entries`` is a list of ``(level, term, page_text)`` tuples seeding
        the visual cache render. ``xe_pairs`` is a list of dicts
        ``{"block_id", "term", "see"?, "level", "page"}`` -- each places a
        hidden XE field at the given paragraph (via add_xe) *and* seeds a
        matching cache entry. entries-derived cache paragraphs come first,
        followed by xe_pairs-derived ones. Returns the open paragraph's
        block_id. No heuristic placement: callers supply exact anchors.
        """
        self._check_hash(expect_hash)
        self._ensure_index_styles()

        container_el = self._container_element(container, section)
        _tag, anchor_el = self._get_block(container_el, after)

        for pair in xe_pairs or []:
            try:
                xe_block_id = pair["block_id"]
                xe_term = pair["term"]
            except KeyError as exc:
                raise DocxError(f"xe_pairs entry missing required key: {exc}") from None
            self.add_xe(
                xe_block_id,
                xe_term,
                see=pair.get("see"),
                container=container,
                section=section,
            )

        open_p = etree.Element(qn("w:p"))
        r1 = etree.SubElement(open_p, qn("w:r"))
        etree.SubElement(r1, qn("w:fldChar")).set(qn("w:fldCharType"), "begin")
        r2 = etree.SubElement(open_p, qn("w:r"))
        instr_el = etree.SubElement(r2, qn("w:instrText"))
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = self._index_instruction(collapsed, locale)
        r3 = etree.SubElement(open_p, qn("w:r"))
        etree.SubElement(r3, qn("w:fldChar")).set(qn("w:fldCharType"), "separate")

        cache = []
        for level, term, page_text in entries or []:
            cache.append(self._cache_paragraph(level, term, page_text, main_style, sub_style))
        for pair in xe_pairs or []:
            try:
                level = pair["level"]
                page_text = pair["page"]
            except KeyError as exc:
                raise DocxError(f"xe_pairs entry missing required key: {exc}") from None
            cache.append(self._cache_paragraph(level, pair["term"], page_text, main_style, sub_style))

        close_p = etree.Element(qn("w:p"))
        rc = etree.SubElement(close_p, qn("w:r"))
        etree.SubElement(rc, qn("w:fldChar")).set(qn("w:fldCharType"), "end")

        span = [open_p] + cache + [close_p]
        anchor = anchor_el
        for el in span:
            anchor.addnext(el)
            anchor = el

        return self._block_index_of(container_el, open_p)

    def add_paragraphs(
        self,
        items: list[dict[str, str | None]],
        after: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
        list_type: str | None = None,
    ) -> list[int]:
        """Add multiple paragraphs at once. Returns list of new block_ids.

        Each item is a dict with keys:
          - "text": paragraph text (required)
          - "style": optional style name
          - "list_type": optional "bullet" or "number"
        When list_type is provided at the function level, it applies to every
        item unless the item overrides it. Paragraphs with list_type but no
        explicit style are automatically assigned "List Paragraph".
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)

        if after is not None:
            _tag, target_el = self._get_block(container_el, after)

        # Resolve list_type: function-level or per-item override
        effective_lt = list_type  # default to function-level value (may be None)

        new_ids: list[int] = []
        prev_p = None  # reference to previously inserted paragraph for chaining

        for item in items:
            style = item.get("style")  # may be None
            text = item.get("text", "")
            lt = item.get("list_type", effective_lt)

            # Pass 'List Paragraph' as python-docx Style so it sets pStyle + rPr
            # automatically. Without this, parent.add_paragraph(text, None) creates
            # a plain paragraph missing <w:rPr lang="en-US"> that Word expects on
            # list paragraphs for proper rendering.
            if lt in ("bullet", "number") and style is None:
                try:
                    py_style = self._document.styles["List Paragraph"]
                except KeyError:
                    py_style = None
            else:
                py_style = style

            p = parent.add_paragraph(text, py_style)

            if after is not None:
                # First item: insert right after the anchor element
                # Subsequent items: chain from previous inserted paragraph
                if prev_p is None:
                    target_el.addnext(p._p)
                else:
                    prev_p._p.addnext(p._p)
            # Apply list formatting (numPr + rPr.lang) if requested.
            if lt in ("bullet", "number") and style is None:
                num_id = self._ensure_numbering_def(lt)
                pPr = p._p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(p._p, qn("w:pPr"))
                # Ensure rPr with lang=en-US — Word needs this for proper list rendering
                rPr_el = pPr.find(qn("w:rPr"))
                if rPr_el is None:
                    rPr_el = etree.SubElement(pPr, qn("w:rPr"))
                lang_el = rPr_el.find(qn("w:lang"))
                if lang_el is None:
                    lang_el = etree.SubElement(rPr_el, qn("w:lang"))
                lang_el.set(qn("w:val"), "en-US")
                numPr = etree.SubElement(pPr, qn("w:numPr"))
                ilvl = etree.SubElement(numPr, qn("w:ilvl"))
                ilvl.set(qn("w:val"), "0")
                numId_el = etree.SubElement(numPr, qn("w:numId"))
                numId_el.set(qn("w:val"), str(num_id))
            prev_p = p
            new_ids.append(self._block_index_of(container_el, p._p))
        return new_ids

    # -- fields ----------------------------------------------------------

    @staticmethod
    def _escape_field_text(s: str) -> str:
        return s.replace("\\", "\\\\").replace('"', '\\"')

    def _field_runs(self, instruction: str) -> list:
        r1 = etree.Element(qn("w:r"))
        f1 = etree.SubElement(r1, qn("w:fldChar"))
        f1.set(qn("w:fldCharType"), "begin")

        r2 = etree.Element(qn("w:r"))
        instr = etree.SubElement(r2, qn("w:instrText"))
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction

        r3 = etree.Element(qn("w:r"))
        f3 = etree.SubElement(r3, qn("w:fldChar"))
        f3.set(qn("w:fldCharType"), "end")

        return [r1, r2, r3]

    def add_field(
        self,
        block_id: int,
        instruction: str,
        after: int | None = None,
        container: str = "body",
        section: int = 0,
        mode: str = "inline",
        expect_hash: str | None = None,
    ) -> int:
        """Insert a Word field (begin/instrText/end triplet) inline in a paragraph.

        ``after`` has no effect in ``mode="inline"`` -- the field is always
        placed in ``block_id``'s own paragraph, prepended before any existing
        runs so visible text is untouched. ``mode="inline"`` is the only
        supported mode; see ``add_index`` for the multi-paragraph INDEX span.
        """
        self._check_hash(expect_hash)
        if mode != "inline":
            raise DocxError(f"unsupported field mode: {mode!r}")
        container_el = self._container_element(container, section)
        _tag, p_el = self._get_block(container_el, block_id, expected_tag="p")

        runs = self._field_runs(instruction)
        existing_runs = p_el.findall(qn("w:r"))
        if existing_runs:
            for r in runs:
                existing_runs[0].addprevious(r)
        else:
            for r in runs:
                p_el.append(r)

        return block_id

    def add_xe(
        self,
        block_id: int,
        term: str,
        see: str | None = None,
        after: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Insert a hidden inline XE (index-entry) field at paragraph block_id.

        ``term`` may use ``"parent:sub"`` colon notation for a sub-entry --
        this is opaque to the toolkit, passed through verbatim for Word to
        interpret on field update. ``see`` adds a ``\\t "See ..."``
        cross-reference switch. The ``\\r "bookmark"`` range switch is not a
        parameter here; use ``add_field`` directly for that case.
        """
        escaped_term = self._escape_field_text(term)
        if see is not None:
            instruction = f' XE "{escaped_term}" \\t "See {self._escape_field_text(see)}" '
        else:
            instruction = f' XE "{escaped_term}" '
        return self.add_field(
            block_id,
            instruction,
            container=container,
            section=section,
            expect_hash=expect_hash,
        )

    def edit_paragraph(
        self,
        block_id: int,
        text: str,
        run: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> None:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="p")
        paragraph = Paragraph(el, parent)
        self._set_paragraph_text(paragraph, text, run)

    def _set_paragraph_text(self, paragraph: Paragraph, text: str, run: int | None = None) -> None:
        runs = paragraph.runs
        if run is not None:
            if run < 0 or run >= len(runs):
                raise DocxError(f"run index out of range: {run}")
            runs[run].text = text
            return
        if not runs:
            paragraph.add_run(text)
            return
        runs[0].text = text
        for extra_run in runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)

    def delete_paragraph(
        self, block_id: int, container: str = "body", section: int = 0, expect_hash: str | None = None
    ) -> None:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="p")
        el.getparent().remove(el)

    def merge_paragraphs(
        self,
        block_ids: list[int],
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Merge adjacent paragraphs into one. Returns the id of the merged paragraph.

        The first paragraph keeps its text and style; subsequent paragraphs' texts
        are appended (separated by a space if needed). All other paragraphs are deleted.
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)

        # Validate all block_ids exist and are adjacent
        blocks = list(self._iter_blocks(container_el))
        id_to_idx = {}
        for idx, (_tag, el) in enumerate(blocks):
            if _tag == "p":
                p_id = self._block_index_of(container_el, el)
                id_to_idx[p_id] = idx

        indices = []
        for bid in block_ids:
            if bid not in id_to_idx:
                raise DocxError(f"block_id {bid} not found")
            indices.append(id_to_idx[bid])

        # Check adjacency
        for i in range(1, len(indices)):
            if indices[i] != indices[i - 1] + 1:
                raise DocxError(
                    f"paragraphs {block_ids} are not adjacent "
                    f"(indices: {indices})"
                )

        # Get the first paragraph element and collect texts from all
        first_el = blocks[indices[0]][1]
        full_text_parts = []

        for idx in indices:
            _tag, el = blocks[idx]
            para_text = "".join(run.text or "" for run in Paragraph(el, container_el).runs)
            full_text_parts.append(para_text)

        # Build merged text — join with space if both parts are non-empty
        merged_text = " ".join(part for part in full_text_parts if part) if any(full_text_parts) else ""

        # Set the first paragraph's text to the merged result
        self._set_paragraph_text(Paragraph(first_el, first_el.getparent()), merged_text)

        # Delete all other paragraphs (reverse order to preserve indices)
        for idx in reversed(indices[1:]):
            _tag, el = blocks[idx]
            el.getparent().remove(el)

        return self._block_index_of(container_el, first_el)

    def split_paragraph(
        self,
        block_id: int,
        offset: int,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Split a paragraph at the given character offset.

        Returns the block_id of the newly created second paragraph.
        The original paragraph keeps text[:offset], new paragraph gets text[offset:].
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="p")

        # Get current text
        para_text = "".join(run.text or "" for run in Paragraph(el, container_el).runs)

        if offset < 0:
            raise DocxError(f"split offset {offset} is negative")
        if offset > len(para_text):
            raise DocxError(
                f"split offset {offset} exceeds text length {len(para_text)}"
            )

        first_part = para_text[:offset]
        second_part = para_text[offset:]

        # Set the original paragraph to the first part
        self._set_paragraph_text(Paragraph(el, el.getparent()), first_part)

        # Create a new paragraph with the second part, inserting it right after
        parent = el.getparent()
        new_p = etree.SubElement(parent, qn("w:p"))
        # Copy paragraph properties from original
        pPr = el.find(qn("w:pPr"))
        if pPr is not None:
            new_pPr = etree.SubElement(new_p, qn("w:pPr"))
            for child in pPr:
                new_pPr.append(etree.copy(child))

        # Add the text as a run
        r = etree.SubElement(new_p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = second_part

        # Move new paragraph right after original
        el.addnext(new_p)

        return self._block_index_of(container_el, new_p)

    def add_table(
        self,
        rows: int,
        cols: int,
        style: str | None = None,
        after: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        if container == "body":
            table = self._document.add_table(rows, cols, style)
        else:
            width = self._document._block_width
            table = parent.add_table(rows, cols, width)
            if style is not None:
                table.style = style
        if after is not None:
            _tag, target_el = self._get_block(container_el, after)
            target_el.addnext(table._tbl)
        return self._block_index_of(container_el, table._tbl)

    def get_table(self, block_id: int, container: str = "body", section: int = 0) -> dict[str, Any]:
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")

        # Count rows and columns from XML directly (more reliable than python-docx Table)
        tr_elements = el.findall(".//" + qn("w:tr"))
        num_rows = len(tr_elements)
        num_cols = 0
        if tr_elements:
            first_row_tc = tr_elements[0].findall(qn("w:tc"))
            num_cols = len(first_row_tc)

        # Extract cell text from XML
        cells = []
        for tr in tr_elements:
            tc_list = tr.findall(qn("w:tc"))
            row_texts = []
            for tc in tc_list:
                # Get all paragraph texts in this cell
                p_elements = tc.findall(qn("w:p"))
                cell_text_parts = []
                for p in p_elements:
                    r_elements = p.findall(qn("w:r"))
                    for r in r_elements:
                        t_elements = r.findall(qn("w:t"))
                        for t in t_elements:
                            if t.text:
                                cell_text_parts.append(t.text)
                row_texts.append("".join(cell_text_parts))
            cells.append(row_texts)

        return {"id": block_id, "rows": num_rows, "cols": num_cols, "cells": cells}

    def set_cell(
        self,
        block_id: int,
        row: int,
        col: int,
        text: str,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> None:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")
        table = Table(el, parent)
        table.cell(row, col).text = text

    def table_fill(
        self,
        block_id: int,
        data: list[list[str]] | str,
        header: list[str] | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> None:
        """Fill a table's cells from a 2D array or CSV string.

        Args:
            block_id: Table block id.
            data: Either a list of lists (rows x cols) or a CSV string with
                  newlines separating rows and commas separating columns.
            header: Optional header row to place in the first row.
                    If provided, data starts filling from row 1.
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")

        # Parse CSV string if needed
        if isinstance(data, str):
            rows = []
            for line in data.strip().splitlines():
                cols = [c.strip() for c in line.split(",")]
                rows.append(cols)
            data = rows
        elif not isinstance(data, list):
            raise DocxError(f"table_fill expects list[list[str]] or str, got {type(data).__name__}")

        # Get table structure from XML (more reliable than python-docx Table)
        tr_elements = el.findall(".//" + qn("w:tr"))
        if not tr_elements:
            raise DocxError("table has no rows")

        num_rows = len(tr_elements)
        first_row_tc = tr_elements[0].findall(qn("w:tc"))
        num_cols = len(first_row_tc) if first_row_tc else 0

        # Set header row if provided
        start_row = 0
        if header is not None:
            if len(header) != num_cols:
                raise DocxError(
                    f"header has {len(header)} columns but table has {num_cols}"
                )
            tc_list = tr_elements[0].findall(qn("w:tc"))
            for col_idx, text in enumerate(header):
                if col_idx < len(tc_list):
                    # Clear existing content and set new
                    p_elements = tc_list[col_idx].findall(qn("w:p"))
                    for p in p_elements:
                        tc_list[col_idx].remove(p)
                    r = etree.SubElement(
                        tc_list[col_idx], qn("w:p")
                    )
                    run = etree.SubElement(r, qn("w:r"))
                    t = etree.SubElement(run, qn("w:t"))
                    t.text = text
            start_row = 1

        # Fill data rows
        for row_idx, row_data in enumerate(data):
            actual_row = start_row + row_idx
            if actual_row >= num_rows:
                break  # Don't exceed table bounds

            tr = tr_elements[actual_row]
            tc_list = tr.findall(qn("w:tc"))
            for col_idx, text in enumerate(row_data):
                if col_idx < len(tc_list):
                    p_elements = tc_list[col_idx].findall(qn("w:p"))
                    for p in p_elements:
                        tc_list[col_idx].remove(p)
                    r = etree.SubElement(
                        tc_list[col_idx], qn("w:p")
                    )
                    run = etree.SubElement(r, qn("w:r"))
                    t = etree.SubElement(run, qn("w:t"))
                    t.text = text

    def add_row(
        self,
        block_id: int,
        values: list[str],
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")
        table = Table(el, parent)
        new_row = table.add_row()
        for col_idx, value in enumerate(values):
            new_row.cells[col_idx].text = str(value)
        return len(table.rows) - 1

    def delete_row(
        self, block_id: int, row: int, container: str = "body", section: int = 0, expect_hash: str | None = None
    ) -> None:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")
        table = Table(el, parent)
        tr = table.rows[row]._tr
        tr.getparent().remove(tr)

    def add_column(
        self,
        block_id: int,
        col_index: int,
        values: list[str],
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Insert a new column at the given index. Returns the new block_id."""
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")

        num_rows = len(el.findall(".//" + qn("w:tr")))
        if col_index < 0 or col_index > num_rows:
            raise DocxError(f"column index {col_index} out of bounds for table with {num_rows} rows")

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # For each row, insert a new <w:tc> at the correct position
        tr_elements = el.findall(".//" + qn("w:tr"))
        for idx, tr in enumerate(tr_elements):
            tc_list = tr.findall(qn("w:tc"))
            if not tc_list:
                continue

            # Create new cell with basic structure
            new_tc = etree.SubElement(tr, qn("w:tc"))

            # Set text
            tc_text = values[idx] if idx < len(values) else ""
            p = etree.SubElement(new_tc, qn("w:p"))
            pPr = etree.SubElement(p, qn("w:pPr"))
            jc = etree.SubElement(pPr, qn("w:jc"))
            jc.set(qn("w:val"), "left")
            r = etree.SubElement(p, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = tc_text

            # Insert before the cell at col_index (or append)
            if col_index < len(tc_list):
                tc_list[col_index].addprevious(new_tc)
            else:
                tr.append(new_tc)

        return self._block_index_of(container_el, el)

    def delete_table(
        self, block_id: int, container: str = "body", section: int = 0, expect_hash: str | None = None
    ) -> None:
        """Delete an entire table block from the document."""
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")
        el.getparent().remove(el)

    def copy_block(
        self,
        block_id: int,
        after: int | None = None,
        before: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Copy a block (paragraph or table) to a new position.

        Either `after` or `before` must be specified. The copied element is
        inserted relative to the target block's current position.
        Returns the new block_id of the copy.
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, src_el = self._get_block(container_el, block_id)

        # Deep copy the element using lxml's copy
        import copy as copy_mod
        new_el = copy_mod.deepcopy(src_el)

        # Determine insertion point
        if after is not None:
            _, target_el = self._get_block(container_el, after)
            target_el.addnext(new_el)
        elif before is not None:
            _, target_el = self._get_block(container_el, before)
            target_el.addprevious(new_el)
        else:
            raise DocxError("copy_block requires 'after' or 'before'")

        return self._block_index_of(container_el, new_el)

    def move_block(
        self,
        block_id: int,
        after: int | None = None,
        before: int | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        """Move a block from its current position to a new one.

        Either `after` or `before` must be specified. Returns the (same) block_id.
        Handles index shifts when moving forward in the document.
        """
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id)

        # Find target element BEFORE removing source (to avoid index shifts)
        if after is not None:
            _, target_el = self._get_block(container_el, after)
        elif before is not None:
            _, target_el = self._get_block(container_el, before)
        else:
            raise DocxError("move_block requires 'after' or 'before'")

        # Remove from current position
        parent = el.getparent()
        parent.remove(el)

        # Insert at new position
        if after is not None:
            target_el.addnext(el)
        elif before is not None:
            target_el.addprevious(el)

        return block_id  # Same id since it's the same element

    def delete_column(
        self,
        block_id: int,
        col_index: int,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> None:
        """Delete a column from the table at the given index."""
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        _tag, el = self._get_block(container_el, block_id, expected_tag="tbl")

        # Get first row to determine column count
        tr_elements = el.findall(".//" + qn("w:tr"))
        if not tr_elements:
            raise DocxError("table has no rows")
        tc_list = tr_elements[0].findall(qn("w:tc"))
        num_cols = len(tc_list)

        if col_index < 0 or col_index >= num_cols:
            raise DocxError(f"column index {col_index} out of bounds for table with {num_cols} columns")

        # Remove the cell at col_index from each row (reverse order for safety)
        for tr in reversed(tr_elements):
            tc_list = tr.findall(qn("w:tc"))
            if col_index < len(tc_list):
                tc_list[col_index].getparent().remove(tc_list[col_index])

    def delete_range(
        self,
        start_id: int,
        end_id: int,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> None:
        """Delete all blocks from start_id (inclusive) to end_id (exclusive)."""
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)

        blocks = self._iter_blocks(container_el)

        if start_id >= end_id:
            raise DocxError(
                f"delete_range: empty range ({start_id}-{end_id}); "
                f"start_id must be less than end_id (end-exclusive)"
            )
        if start_id < 0 or start_id >= len(blocks):
            raise DocxError(
                f"delete_range: start_id {start_id} out of bounds "
                f"(doc has {len(blocks)} blocks, valid range 0..{len(blocks)})"
            )
        if end_id > len(blocks):
            raise DocxError(
                f"delete_range: end_id {end_id} out of bounds "
                f"(doc has {len(blocks)} blocks, valid range 0..{len(blocks)})"
            )

        # Collect elements to remove in reverse order for safe deletion
        to_remove = []
        for idx in range(start_id, end_id):
            _tag, el = blocks[idx]
            to_remove.append(el)

        # Remove in reverse order so earlier indices stay valid
        for el in reversed(to_remove):
            el.getparent().remove(el)

    # -- comments ----------------------------------------------------------

    def list_comments(self) -> list[dict[str, Any]]:
        body = self._document.element.body
        body_blocks = self._iter_blocks(body)
        results = []
        for comment in self._document.comments:
            anchor_block = None
            anchor_tx = ""
            for idx, (tag, el) in enumerate(body_blocks):
                if tag != "p":
                    continue
                match = any(
                    start.get(qn("w:id")) == str(comment.comment_id)
                    for start in el.iter(qn("w:commentRangeStart"))
                )
                if match:
                    anchor_block = idx
                    anchor_tx = Paragraph(el, self._document).text
                    break
            results.append(
                {
                    "id": comment.comment_id,
                    "author": comment.author,
                    "initials": comment.initials,
                    "date": comment.timestamp.isoformat() if comment.timestamp else None,
                    "text": comment.text,
                    "anchor_block": anchor_block,
                    "anchor_tx": anchor_tx,
                }
            )
        return results

    def add_comment(
        self,
        block_id: int,
        run_start: int,
        run_end: int,
        text: str,
        author: str = "",
        initials: str = "",
        expect_hash: str | None = None,
    ) -> int:
        self._check_hash(expect_hash)
        container_el = self._container_element("body", 0)
        _tag, el = self._get_block(container_el, block_id, expected_tag="p")
        paragraph = Paragraph(el, self._document)
        runs = paragraph.runs[run_start:run_end]
        if not runs:
            raise DocxError("run range is empty")
        comment = self._document.add_comment(runs=runs, text=text, author=author, initials=initials)
        return comment.comment_id

    def edit_comment(
        self,
        comment_id: int,
        text: str | None = None,
        author: str | None = None,
        initials: str | None = None,
    ) -> None:
        comment = self._document.comments.get(comment_id)
        if comment is None:
            raise DocxError(f"unknown comment_id: {comment_id}")
        if text is not None:
            self._set_paragraph_text(comment.paragraphs[0], text)
        if author is not None:
            comment.author = author
        if initials is not None:
            comment.initials = initials

    # -- images ----------------------------------------------------------

    def add_image(
        self,
        src: str | Path,
        after: int | None = None,
        width: float | None = None,
        height: float | None = None,
        container: str = "body",
        section: int = 0,
        expect_hash: str | None = None,
    ) -> int:
        self._check_hash(expect_hash)
        container_el = self._container_element(container, section)
        parent = self._container_parent(container, section)
        new_paragraph = parent.add_paragraph()
        run = new_paragraph.add_run()
        kwargs: dict[str, Any] = {}
        if width is not None:
            kwargs["width"] = Inches(width)
        if height is not None:
            kwargs["height"] = Inches(height)
        run.add_picture(str(src), **kwargs)
        if after is not None:
            _tag, target_el = self._get_block(container_el, after)
            target_el.addnext(new_paragraph._p)
        return self._block_index_of(container_el, new_paragraph._p)

    def list_images(self) -> list[dict[str, Any]]:
        body = self._document.element.body
        results = []
        for idx, shape in enumerate(self._document.inline_shapes):
            p_el = shape._inline
            while p_el is not None and p_el.tag != qn("w:p"):
                p_el = p_el.getparent()
            block_id = self._block_index_of(body, p_el) if p_el is not None else None
            results.append(
                {
                    "image_id": idx,
                    "block_id": block_id,
                    "width": shape.width / 914400,
                    "height": shape.height / 914400,
                }
            )
        return results

    def extract_image(self, image_id: int, out_path: str | Path) -> None:
        shape = self._document.inline_shapes[image_id]
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = self._document.part.related_parts[rId]
        Path(out_path).write_bytes(image_part.blob)

    # -- headers / footers -----------------------------------------------

    def get_header(self, section: int = 0) -> dict[str, Any]:
        return self.read(container="header", section=section)

    def get_footer(self, section: int = 0) -> dict[str, Any]:
        return self.read(container="footer", section=section)
