from docx import Document as RawDocument

from docx_toolkit import DocxDocument


def test_save_and_reopen_preserves_structure(sample_doc: DocxDocument, doc_path):
    sample_doc.save(doc_path)
    reopened = DocxDocument.open(doc_path)

    original_blocks = sample_doc.read()["blocks"]
    reopened_blocks = reopened.read()["blocks"]

    assert len(original_blocks) == len(reopened_blocks)
    for a, b in zip(original_blocks, reopened_blocks):
        assert a["ty"] == b["ty"]
        assert a.get("st") == b.get("st")
        if a["ty"] == "p":
            assert a["tx"] == b["tx"]


def test_editing_does_not_alter_style_definitions(sample_doc: DocxDocument, doc_path):
    sample_doc.save(doc_path)
    before_styles = {s["name"] for s in DocxDocument.open(doc_path).list_styles()}

    doc = DocxDocument.open(doc_path)
    block_id = doc.outline()[0]["id"]
    doc.edit_paragraph(block_id, "Edited Title Text")
    doc.save(doc_path)

    after_styles = {s["name"] for s in DocxDocument.open(doc_path).list_styles()}
    assert before_styles == after_styles


def test_untouched_paragraph_formatting_survives_unrelated_edit(sample_doc: DocxDocument, doc_path):
    sample_doc.save(doc_path)

    raw = RawDocument(str(doc_path))
    raw.paragraphs[2].runs[0].italic = True
    raw.save(str(doc_path))

    doc = DocxDocument.open(doc_path)
    unrelated_id = doc.outline()[0]["id"]
    doc.edit_paragraph(unrelated_id, "Changed only the title")
    doc.save(doc_path)

    final = RawDocument(str(doc_path))
    assert final.paragraphs[2].runs[0].italic is True


def test_full_roundtrip_of_all_content_types(blank_doc: DocxDocument, tmp_path, doc_path):
    import struct
    import zlib

    def make_png():
        def chunk(tag, data):
            return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

        sig = b"\x89PNG\r\n\x1a\n"
        ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        idat = chunk(b"IDAT", zlib.compress(b"\x00" + bytes((0, 255, 0))))
        iend = chunk(b"IEND", b"")
        return sig + ihdr + idat + iend

    png_path = tmp_path / "img.png"
    png_path.write_bytes(make_png())

    p1 = blank_doc.add_paragraph("Intro paragraph")
    tbl = blank_doc.add_table(rows=1, cols=2, after=p1)
    blank_doc.set_cell(tbl, 0, 0, "Key")
    blank_doc.set_cell(tbl, 0, 1, "Value")
    img_block = blank_doc.add_image(png_path, after=tbl)
    blank_doc.add_comment(p1, 0, 1, "a comment", author="tester")
    blank_doc.add_paragraph("Header", container="header", section=0)
    blank_doc.add_paragraph("Footer", container="footer", section=0)

    blank_doc.save(doc_path)
    reopened = DocxDocument.open(doc_path)

    assert len(reopened.list_comments()) == 1
    assert len(reopened.list_images()) == 1
    assert reopened.get_table(tbl)["cells"] == [["Key", "Value"]]
    assert "Header" in [b["tx"] for b in reopened.get_header(0)["blocks"] if b["ty"] == "p"]
    assert "Footer" in [b["tx"] for b in reopened.get_footer(0)["blocks"] if b["ty"] == "p"]
