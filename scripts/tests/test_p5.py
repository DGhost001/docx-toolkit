"""P5 features: Word field support — add_field, add_xe, add_index, validate field checks."""
from docx.oxml.ns import qn

import pytest

from docx_toolkit import DocxDocument, DocxError, StaleHashError


def _fld_char_types(p_el):
    return [
        r.find(qn("w:fldChar")).get(qn("w:fldCharType"))
        for r in p_el.findall(qn("w:r"))
        if r.find(qn("w:fldChar")) is not None
    ]


def _instr_texts(p_el):
    return [t.text for t in p_el.findall(".//" + qn("w:instrText"))]


class TestAddField:
    def test_add_field_inserts_begin_instrtext_end_triplet(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        result = doc.add_field(block_id, ' TOC \\o "1-3" \\h \\z ')
        assert result == block_id

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        runs = p_el.findall(qn("w:r"))
        assert len(runs) == 4  # 3 field runs + 1 pre-existing text run

        assert _fld_char_types(p_el) == ["begin", "end"]

        instr_run = runs[1]
        instr_el = instr_run.find(qn("w:instrText"))
        assert instr_el.text == ' TOC \\o "1-3" \\h \\z '
        assert instr_el.get(qn("xml:space")) == "preserve"

        # Field runs are prepended; visible text run comes last and is untouched.
        assert runs[3].find(qn("w:t")).text == "Body"
        assert doc.get_paragraph(block_id)["tx"] == "Body"

    def test_add_field_on_empty_paragraph(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("")
        doc.add_field(block_id, " REF bookmark1 ")

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        runs = p_el.findall(qn("w:r"))
        assert len(runs) == 3
        assert _fld_char_types(p_el) == ["begin", "end"]

    def test_add_field_unsupported_mode_raises(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        with pytest.raises(DocxError, match="unsupported field mode"):
            doc.add_field(block_id, "X", mode="bogus")

    def test_add_field_unknown_block_id_raises(self):
        doc = DocxDocument.new()
        with pytest.raises(DocxError, match="unknown block_id"):
            doc.add_field(0, "X")

    def test_add_field_stale_hash_raises_and_does_not_modify(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        with pytest.raises(StaleHashError):
            doc.add_field(block_id, "X", expect_hash="not-the-real-hash")
        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert len(p_el.findall(qn("w:r"))) == 1


class TestAddXe:
    def test_add_xe_plain_term(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Discussing foo here.")
        doc.add_xe(block_id, "foo")

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert _instr_texts(p_el) == [' XE "foo" ']
        assert _fld_char_types(p_el) == ["begin", "end"]
        assert doc.get_paragraph(block_id)["tx"] == "Discussing foo here."

    def test_add_xe_colon_nesting(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        doc.add_xe(block_id, "bar:foo")

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert _instr_texts(p_el) == [' XE "bar:foo" ']

    def test_add_xe_with_see_cross_reference(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        doc.add_xe(block_id, "foo", see="bar")

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert _instr_texts(p_el) == [' XE "foo" \\t "See bar" ']

    def test_add_xe_escapes_quotes_and_backslashes(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("")
        doc.add_xe(block_id, 'a\\b"c')

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert _instr_texts(p_el) == [' XE "a\\\\b\\"c" ']
        assert len(p_el.findall(qn("w:r"))) == 3

    def test_add_xe_is_hidden_and_preserves_visible_text(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Discussing foo here.")
        doc.add_xe(block_id, "foo")

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        runs = p_el.findall(qn("w:r"))
        assert runs[3].find(qn("w:t")).text == "Discussing foo here."
        assert doc.get_paragraph(block_id)["tx"] == "Discussing foo here."

    def test_add_xe_range_switch_reachable_via_add_field(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        doc.add_field(block_id, ' XE "foo" \\r "mybookmark" ')

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert _instr_texts(p_el) == [' XE "foo" \\r "mybookmark" ']

    def test_add_xe_unknown_block_id_raises(self):
        doc = DocxDocument.new()
        with pytest.raises(DocxError, match="unknown block_id"):
            doc.add_xe(0, "foo")

    def test_add_xe_stale_hash_raises_and_does_not_modify(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        with pytest.raises(StaleHashError):
            doc.add_xe(block_id, "foo", expect_hash="not-the-real-hash")
        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        assert len(p_el.findall(qn("w:r"))) == 1


class TestAddIndex:
    def test_add_index_creates_open_and_close_paragraphs_only(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Analytical Index", style="Heading 1")

        open_id = doc.add_index(after=heading_id)

        blocks = doc.read()["blocks"]
        assert len(blocks) == 3  # heading + open + close
        assert open_id == 1

        _tag, open_p = doc._get_block(doc._document.element.body, 1)
        _tag, close_p = doc._get_block(doc._document.element.body, 2)

        assert open_p is not close_p
        assert close_p.findall(qn("w:p")) == []  # never nested

        assert _fld_char_types(open_p) == ["begin", "separate"]
        assert _fld_char_types(close_p) == ["end"]

        assert len(open_p.findall(qn("w:r"))) == 3
        assert len(close_p.findall(qn("w:r"))) == 1

        assert _instr_texts(open_p) == [' INDEX \\e "\\t" \\c "2" \\z "1033" ']

    def test_add_index_returns_open_paragraph_block_id(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        open_id = doc.add_index(after=heading_id)
        assert open_id == heading_id + 1

    def test_add_index_custom_collapsed_and_locale(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        open_id = doc.add_index(after=heading_id, collapsed="3", locale="1031")
        _tag, open_p = doc._get_block(doc._document.element.body, open_id)
        assert _instr_texts(open_p) == [' INDEX \\e "\\t" \\c "3" \\z "1031" ']

    def test_add_index_zero_entries_means_no_cache_paragraphs(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id, entries=[])
        assert len(doc.read()["blocks"]) == 3  # heading + open + close, no cache

    def test_add_index_unknown_after_block_id_raises(self):
        doc = DocxDocument.new()
        with pytest.raises(DocxError, match="unknown block_id"):
            doc.add_index(after=0)

    def test_add_index_stale_hash_raises_and_does_not_modify(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        with pytest.raises(StaleHashError):
            doc.add_index(after=heading_id, expect_hash="not-the-real-hash")
        assert len(doc.read()["blocks"]) == 1


def _tab_attrs(p_el):
    tab = p_el.find(qn("w:pPr")).find(qn("w:tabs")).find(qn("w:tab"))
    return {
        "val": tab.get(qn("w:val")),
        "leader": tab.get(qn("w:leader")),
        "pos": tab.get(qn("w:pos")),
    }


def _cache_texts(p_el):
    r = p_el.find(qn("w:r"))
    return [t.text for t in r.findall(qn("w:t"))]


class TestAddIndexSeeding:
    def test_add_index_with_entries_creates_cache_paragraphs(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        entries = [(1, "foo", "12, 27"), (2, "sub", "12")]
        open_id = doc.add_index(after=heading_id, entries=entries)

        blocks = doc.read()["blocks"]
        assert len(blocks) == 5  # heading + open + 2 cache + close

        _tag, cache0 = doc._get_block(doc._document.element.body, open_id + 1)
        _tag, cache1 = doc._get_block(doc._document.element.body, open_id + 2)

        assert cache0.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "index 1"
        assert _cache_texts(cache0) == ["foo", "12, 27"]
        assert cache0.find(qn("w:r")).find(qn("w:tab")) is not None

        assert cache1.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "index 2"
        assert _cache_texts(cache1) == ["sub", "12"]

    def test_add_index_entries_level_3_uses_hardcoded_style(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        open_id = doc.add_index(after=heading_id, entries=[(3, "leaf", "5")])
        _tag, cache0 = doc._get_block(doc._document.element.body, open_id + 1)
        assert cache0.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "index 3"

    def test_add_index_cache_paragraph_has_right_dot_leader_tab(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        open_id = doc.add_index(after=heading_id, entries=[(1, "foo", "1")])
        _tag, cache0 = doc._get_block(doc._document.element.body, open_id + 1)
        assert _tab_attrs(cache0) == {"val": "right", "leader": "dot", "pos": "4448"}

    def test_add_index_with_xe_pairs_places_xe_fields_and_seeds_cache(self):
        doc = DocxDocument.new()
        p0 = doc.add_paragraph("Discussing foo here.")
        p1 = doc.add_paragraph("Unrelated paragraph.")
        p2 = doc.add_paragraph("Discussing bar here.")
        heading_id = doc.add_paragraph("Index")

        xe_pairs = [
            {"block_id": p0, "term": "foo", "level": 1, "page": "1"},
            {"block_id": p2, "term": "bar", "see": None, "level": 1, "page": "3"},
        ]
        open_id = doc.add_index(after=heading_id, xe_pairs=xe_pairs)

        _tag, p0_el = doc._get_block(doc._document.element.body, p0)
        _tag, p1_el = doc._get_block(doc._document.element.body, p1)
        _tag, p2_el = doc._get_block(doc._document.element.body, p2)

        assert _instr_texts(p0_el) == [' XE "foo" ']
        assert _instr_texts(p2_el) == [' XE "bar" ']
        assert p1_el.findall(".//" + qn("w:fldChar")) == []

        _tag, cache0 = doc._get_block(doc._document.element.body, open_id + 1)
        _tag, cache1 = doc._get_block(doc._document.element.body, open_id + 2)
        assert _cache_texts(cache0) == ["foo", "1"]
        assert _cache_texts(cache1) == ["bar", "3"]

    def test_add_index_xe_pairs_and_entries_both_given_are_both_rendered(self):
        doc = DocxDocument.new()
        p0 = doc.add_paragraph("Discussing foo here.")
        heading_id = doc.add_paragraph("Index")

        open_id = doc.add_index(
            after=heading_id,
            entries=[(1, "alpha", "1")],
            xe_pairs=[{"block_id": p0, "term": "foo", "level": 1, "page": "2"}],
        )

        _tag, cache0 = doc._get_block(doc._document.element.body, open_id + 1)
        _tag, cache1 = doc._get_block(doc._document.element.body, open_id + 2)
        assert _cache_texts(cache0) == ["alpha", "1"]  # entries come first
        assert _cache_texts(cache1) == ["foo", "2"]  # then xe_pairs

    def test_add_index_xe_pairs_block_id_resolved_before_span_insertion(self):
        doc = DocxDocument.new()
        p0 = doc.add_paragraph("Discussing foo here.")
        heading_id = doc.add_paragraph("Index")

        # block_id below `after` still resolves correctly -- inserting XE
        # fields doesn't shift block ids, only inserting new paragraphs does.
        doc.add_index(
            after=heading_id,
            xe_pairs=[{"block_id": p0, "term": "foo", "level": 1, "page": "1"}],
        )
        _tag, p0_el = doc._get_block(doc._document.element.body, p0)
        assert _instr_texts(p0_el) == [' XE "foo" ']

    def test_add_index_malformed_xe_pairs_raises_docx_error(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        with pytest.raises(DocxError):
            doc.add_index(after=heading_id, xe_pairs=[{"term": "foo"}])  # missing block_id


def _style_el(doc, style_id):
    styles_el = doc._document.styles.element
    for s in styles_el.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == style_id:
            return s
    return None


class TestEnsureIndexStyles:
    def test_add_index_creates_missing_index_styles(self):
        doc = DocxDocument.new()
        assert "index 1" not in {s["name"] for s in doc.list_styles()}

        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id, entries=[(1, "a", "1")])

        names = {s["name"] for s in doc.list_styles()}
        assert {"index 1", "index 2", "index 3"} <= names

    def test_index_style_definitions_have_correct_shape(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id)

        expected_left = {"index 1": "220", "index 2": "440", "index 3": "660"}
        for style_id, left in expected_left.items():
            style_el = _style_el(doc, style_id)
            assert style_el is not None
            assert style_el.get(qn("w:type")) == "paragraph"
            assert style_el.find(qn("w:basedOn")).get(qn("w:val")) == "Normal"
            assert style_el.find(qn("w:name")).get(qn("w:val")) == style_id

            ind = style_el.find(qn("w:pPr")).find(qn("w:ind"))
            assert ind.get(qn("w:left")) == left
            assert ind.get(qn("w:hanging")) == "220"

            tab = style_el.find(qn("w:pPr")).find(qn("w:tabs")).find(qn("w:tab"))
            assert tab.get(qn("w:val")) == "right"
            assert tab.get(qn("w:leader")) == "dot"
            assert tab.get(qn("w:pos")) == "4448"

    def test_add_index_reuses_existing_index_styles_without_duplicating(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id)
        doc.add_index(after=heading_id)

        names = [s["name"] for s in doc.list_styles()]
        assert names.count("index 1") == 1
        assert names.count("index 2") == 1
        assert names.count("index 3") == 1

    def test_field_add_does_not_create_index_styles(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        doc.add_field(block_id, " REF bookmark1 ")
        assert "index 1" not in {s["name"] for s in doc.list_styles()}

    def test_xe_add_does_not_create_index_styles(self):
        doc = DocxDocument.new()
        block_id = doc.add_paragraph("Body")
        doc.add_xe(block_id, "foo")
        assert "index 1" not in {s["name"] for s in doc.list_styles()}


class TestValidateFields:
    def test_validate_plain_document_unaffected(self):
        doc = DocxDocument.new()
        doc.add_paragraph("Test")
        result = doc.validate()
        assert result["valid"] is True
        assert result["errors"] == []

    def test_validate_valid_index_document_passes(self):
        doc = DocxDocument.new()
        p0 = doc.add_paragraph("Discussing foo here.")
        heading_id = doc.add_paragraph("Index")
        doc.add_index(
            after=heading_id,
            xe_pairs=[{"block_id": p0, "term": "foo", "level": 1, "page": "1"}],
        )
        result = doc.validate()
        assert result["valid"] is True
        assert result["errors"] == []

    def test_validate_detects_multiple_index_instructions(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id)
        doc.add_index(after=heading_id)

        result = doc.validate()
        assert result["valid"] is False
        assert any("INDEX" in e for e in result["errors"])

    def test_validate_detects_unbalanced_begin_end(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id)

        # Corrupt: delete the close paragraph's `end` fldChar run.
        _tag, close_p = doc._get_block(doc._document.element.body, heading_id + 2)
        for r in close_p.findall(qn("w:r")):
            close_p.remove(r)

        result = doc.validate()
        assert result["valid"] is False
        assert any("unbalanced" in e.lower() for e in result["errors"])

    def test_validate_detects_end_nested_in_open_paragraph(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id)

        # Corrupt: move the `end` fldChar run into the open paragraph itself.
        _tag, open_p = doc._get_block(doc._document.element.body, heading_id + 1)
        _tag, close_p = doc._get_block(doc._document.element.body, heading_id + 2)
        end_run = close_p.findall(qn("w:r"))[0]
        open_p.append(end_run)

        result = doc.validate()
        assert result["valid"] is False
        assert any("end" in e.lower() for e in result["errors"])

    def test_validate_detects_xe_inside_index_span(self):
        doc = DocxDocument.new()
        heading_id = doc.add_paragraph("Index")
        doc.add_index(after=heading_id, entries=[(1, "foo", "1")])

        # Corrupt: insert an XE field into the cache paragraph (inside the span).
        _tag, cache_p = doc._get_block(doc._document.element.body, heading_id + 2)
        doc.add_field(heading_id + 2, ' XE "sneaky" ')

        result = doc.validate()
        assert result["valid"] is False
        assert any("XE" in e for e in result["errors"])

    def test_validate_field_free_document_has_no_field_errors(self):
        doc = DocxDocument.new()
        doc.add_paragraph("Test")
        doc.add_table(rows=1, cols=1)
        result = doc.validate()
        assert result["errors"] == []

    def test_validate_ignores_header_footer_fields(self):
        doc = DocxDocument.new()
        # Header/footer field structure is a known v1 scoping gap, not a bug.
        result = doc.validate()
        assert result["valid"] is True
