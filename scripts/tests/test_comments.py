from docx_toolkit import DocxDocument


def test_add_comment_anchored_to_paragraph_runs(blank_doc: DocxDocument, doc_path):
    from docx import Document

    block_id = blank_doc.add_paragraph("Hello, world!")
    comment_id = blank_doc.add_comment(
        block_id, run_start=0, run_end=1, text="I have this to say", author="Agent", initials="AG"
    )
    assert isinstance(comment_id, int)

    blank_doc.save(doc_path)
    reopened = DocxDocument.open(doc_path)
    comments = reopened.list_comments()
    assert len(comments) == 1
    c = comments[0]
    assert c["text"] == "I have this to say"
    assert c["author"] == "Agent"
    assert c["initials"] == "AG"
    assert c["anchor_block"] == block_id
    assert "Hello" in c["anchor_tx"]


def test_edit_comment_text_and_author(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("Some text")
    comment_id = blank_doc.add_comment(block_id, 0, 1, "Original comment", author="A")

    blank_doc.edit_comment(comment_id, text="Updated comment", author="B")

    comments = blank_doc.list_comments()
    updated = next(c for c in comments if c["id"] == comment_id)
    assert updated["text"] == "Updated comment"
    assert updated["author"] == "B"


def test_list_comments_empty_by_default(blank_doc: DocxDocument):
    blank_doc.add_paragraph("No comments here")
    assert blank_doc.list_comments() == []


def test_list_comments_skips_table_blocks_when_locating_anchor(blank_doc: DocxDocument):
    blank_doc.add_table(rows=1, cols=1)
    block_id = blank_doc.add_paragraph("Paragraph after table")
    comment_id = blank_doc.add_comment(block_id, 0, 1, "anchored past a table")

    comment = next(c for c in blank_doc.list_comments() if c["id"] == comment_id)
    assert comment["anchor_block"] == block_id


def test_multiple_comments_on_different_paragraphs(blank_doc: DocxDocument):
    b1 = blank_doc.add_paragraph("First paragraph")
    b2 = blank_doc.add_paragraph("Second paragraph")

    c1 = blank_doc.add_comment(b1, 0, 1, "comment one")
    c2 = blank_doc.add_comment(b2, 0, 1, "comment two")

    comments = blank_doc.list_comments()
    assert {c["id"] for c in comments} == {c1, c2}
