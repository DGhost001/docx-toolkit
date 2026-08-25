import pytest

from docx_toolkit import DocxDocument, StaleHashError


def test_add_paragraph_returns_block_id(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("Hello, world!", style="Normal")
    assert isinstance(block_id, int)

    para = blank_doc.get_paragraph(block_id)
    assert para["tx"] == "Hello, world!"
    assert para["st"] == "Normal"


def test_add_paragraph_after_specific_block(blank_doc: DocxDocument):
    first = blank_doc.add_paragraph("First")
    second = blank_doc.add_paragraph("Second")
    inserted = blank_doc.add_paragraph("Inserted", after=first)

    blocks = blank_doc.read()["blocks"]
    texts = [b["tx"] for b in blocks if b["ty"] == "p"]
    assert texts == ["First", "Inserted", "Second"]
    # block_id is positional (see core.py docstring), so "inserted" landing on
    # the position vacated by the reordering is expected, not a collision.
    assert blocks[inserted]["tx"] == "Inserted"


def test_edit_paragraph_replaces_text(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("Original text")
    blank_doc.edit_paragraph(block_id, "Updated text")
    assert blank_doc.get_paragraph(block_id)["tx"] == "Updated text"


def test_edit_paragraph_preserves_first_run_formatting(blank_doc: DocxDocument, doc_path):
    block_id = blank_doc.add_paragraph("bold text")
    from docx import Document

    blank_doc.save(doc_path)
    raw = Document(str(doc_path))
    raw.paragraphs[-1].runs[0].bold = True
    raw.save(str(doc_path))

    reopened = DocxDocument.open(doc_path)
    reopened.edit_paragraph(block_id, "bold text updated")
    reopened.save(doc_path)

    final = Document(str(doc_path))
    run = final.paragraphs[-1].runs[0]
    assert run.text == "bold text updated"
    assert run.bold is True


def test_edit_specific_run_only(blank_doc: DocxDocument, doc_path):
    from docx import Document

    block_id = blank_doc.add_paragraph("")
    blank_doc.save(doc_path)

    raw = Document(str(doc_path))
    p = raw.paragraphs[-1]
    p.add_run("Hello ")
    p.add_run("World")
    raw.save(str(doc_path))

    doc = DocxDocument.open(doc_path)
    doc.edit_paragraph(block_id, "Universe", run=1)
    para = doc.get_paragraph(block_id)
    assert para["tx"] == "Hello Universe"


def test_delete_paragraph(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("To be removed")
    other = blank_doc.add_paragraph("Keep me")
    blank_doc.delete_paragraph(block_id)

    blocks = blank_doc.read()["blocks"]
    texts = [b["tx"] for b in blocks if b["ty"] == "p"]
    assert "To be removed" not in texts
    assert "Keep me" in texts


def test_edit_paragraph_with_correct_hash_succeeds(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("Text")
    current_hash = blank_doc.content_hash
    blank_doc.edit_paragraph(block_id, "New text", expect_hash=current_hash)
    assert blank_doc.get_paragraph(block_id)["tx"] == "New text"


def test_edit_paragraph_with_stale_hash_raises_and_does_not_modify(blank_doc: DocxDocument):
    block_id = blank_doc.add_paragraph("Text")
    stale_hash = "not-the-real-hash"
    with pytest.raises(StaleHashError):
        blank_doc.edit_paragraph(block_id, "Should not apply", expect_hash=stale_hash)
    assert blank_doc.get_paragraph(block_id)["tx"] == "Text"


def test_get_paragraph_unknown_block_id_raises(blank_doc: DocxDocument):
    from docx_toolkit import DocxError

    with pytest.raises(DocxError):
        blank_doc.get_paragraph(9999)
