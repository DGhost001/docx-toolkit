"""P6 features: CLI subcommands for field-add/xe-add/index-add, batch chaining."""
import json

from docx_toolkit import DocxDocument, cli


def _run(*args: str) -> dict:
    return cli.run(list(args))


class TestFieldAddCLI:
    def test_field_add_cli_inserts_triplet(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        added = _run("para-add", path, "Body")
        block_id = added["id"]

        result = _run("field-add", path, str(block_id), "--instr", ' TOC \\o "1-3" ')
        assert "id" in result
        assert "hash" in result

        doc = DocxDocument.open(path)
        from docx.oxml.ns import qn

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        runs = p_el.findall(qn("w:r"))
        assert len(runs) == 4
        assert doc.get_paragraph(block_id)["tx"] == "Body"

    def test_field_add_cli_unknown_block_id_does_not_save(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        result = _run("field-add", path, "0", "--instr", "X")
        assert "error" in result
        assert result["type"] == "DocxError"


class TestXeAddCLI:
    def test_xe_add_cli_hidden_field(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        added = _run("para-add", path, "Discussing foo here.")
        block_id = added["id"]

        result = _run("xe-add", path, str(block_id), "--term", "foo", "--see", "bar")
        assert "id" in result

        doc = DocxDocument.open(path)
        from docx.oxml.ns import qn

        _tag, p_el = doc._get_block(doc._document.element.body, block_id)
        instr = p_el.find(".//" + qn("w:instrText")).text
        assert instr == ' XE "foo" \\t "See bar" '
        assert doc.get_paragraph(block_id)["tx"] == "Discussing foo here."


class TestIndexAddCLI:
    def test_index_add_cli_builds_span(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        heading = _run("para-add", path, "Index")

        result = _run(
            "index-add", path, "--after", str(heading["id"]),
            "--entries", json.dumps([[1, "foo", "12"]]),
        )
        assert "id" in result

        doc = DocxDocument.open(path)
        blocks = doc.read()["blocks"]
        assert len(blocks) == 4  # heading + open + cache + close
        assert doc.validate()["valid"] is True

    def test_index_add_cli_bad_json_entries_errors(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        heading = _run("para-add", path, "Index")

        result = _run("index-add", path, "--after", str(heading["id"]), "--entries", "not json")
        assert "error" in result


class TestFieldBatchOps:
    def test_batch_xe_add_then_index_add_with_prev_chaining(self):
        from docx_toolkit.cli import _run_batch

        doc = DocxDocument.new()
        ops = [
            {"op": "add_paragraph", "kwargs": {"text": "foo body"}, "as": "p1"},
            {"op": "add_paragraph", "kwargs": {"text": "Index"}, "as": "idx_p"},
            {"op": "add_xe", "kwargs": {"block_id": "$p1", "term": "foo"}, "as": "prev"},
            {
                "op": "add_index",
                "kwargs": {"after": "$idx_p", "entries": [[1, "foo", "1"]]},
                "as": "idx",
            },
        ]
        result = _run_batch(doc, ops)
        assert result["ok"] is True
        assert doc.validate()["valid"] is True

    def test_batch_field_add_unknown_block_reports_structured_error(self):
        from docx_toolkit.cli import _run_batch

        doc = DocxDocument.new()
        ops = [
            {"op": "add_field", "kwargs": {"block_id": 999, "instruction": "X"}},
        ]
        result = _run_batch(doc, ops)
        assert result["ok"] is False
        assert result["failed_at"] == 0
        assert "error" in result

    def test_batch_command_field_add_xe_add_index_add_end_to_end(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _run("new", path)
        ops = [
            {"op": "add_paragraph", "kwargs": {"text": "foo body"}, "as": "p1"},
            {"op": "add_paragraph", "kwargs": {"text": "Index"}, "as": "idx_p"},
            {"op": "add_xe", "kwargs": {"block_id": "$p1", "term": "foo"}},
            {"op": "add_index", "kwargs": {"after": "$idx_p", "entries": [[1, "foo", "1"]]}},
        ]
        result = _run("batch", path, json.dumps(ops))
        assert result["ok"] is True

        doc = DocxDocument.open(path)
        assert doc.validate()["valid"] is True
